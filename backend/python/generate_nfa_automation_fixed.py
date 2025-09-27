import sys
import os
import re
import json
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from datetime import datetime

# Debug Python version
print(f"Python version: {sys.version}", file=sys.stderr)
print(f"Python executable: {sys.executable}", file=sys.stderr)

# ==========================
# Load API Key & Init OpenAI
# ==========================
try:
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("Warning: OPENAI_API_KEY not found", file=sys.stderr)
    else:
        print("OPENAI_API_KEY loaded successfully", file=sys.stderr)
    client = OpenAI(api_key=api_key) if api_key else None
except Exception as e:
    print(f"Error initializing OpenAI client: {e}", file=sys.stderr)
    client = None

# ==========================
# Paths
# ==========================
script_dir = os.path.dirname(os.path.abspath(__file__))
backend_dir = os.path.dirname(script_dir)
uploads_dir = os.path.join(backend_dir, "uploads")

# Fix header image path - look in multiple locations
header_image_path = None
possible_header_paths = [
    os.path.join(uploads_dir, "header.png"),
    os.path.join(backend_dir, "assets", "header.png"),
    os.path.join(backend_dir, "header.png"),
    os.path.join(script_dir, "header.png"),
    os.path.join(backend_dir, "..", "public", "header.png")  # Add public directory
]

for path in possible_header_paths:
    if os.path.exists(path):
        header_image_path = path
        print(f"✅ Found header image at: {path}", file=sys.stderr)
        break

if not header_image_path:
    print("⚠️ Header image not found in any expected location", file=sys.stderr)
    print(f"Searched paths: {possible_header_paths}", file=sys.stderr)

signatures_dir = os.path.join(uploads_dir, "signatures")

# Output directory - FIXED to match server static file serving
env_output_dir = os.getenv("OUTPUT_DIR")
if env_output_dir:
    if env_output_dir.startswith("./backend/"):
        env_output_dir = env_output_dir.replace("./backend/", "./")
    base_output_directory = os.path.abspath(os.path.join(backend_dir, env_output_dir.lstrip("./")))
else:
    # FIXED: Create files directly in backend/generated_letters to match server static serving
    base_output_directory = os.path.join(backend_dir, "generated_letters")

output_directory = os.path.join(base_output_directory, "nfa")
os.makedirs(output_directory, exist_ok=True)

print(f"Final output directory: {output_directory}", file=sys.stderr)

# ==========================
# AI Helper
# ==========================
def generate_ai_nfa_from_summary(subject, summary, nfa_type="reimbursement", need_bullets=False, facts_only=False):
    if not client:
        return f"{subject}\n\nRequest for approval regarding {summary}. The above proposal is submitted for approval."

    # Create ultra-concise prompt for single page following strict template
    prompt = f"""
Create an EXCELLENT professional NFA document following this EXACT structure for SINGLE PAGE output:

USER INPUTS:
Subject: {subject}
Summary: {summary}
Type: {nfa_type}
Need Bullets: {need_bullets}

REQUIRED OUTPUT FORMAT (MUST FOLLOW EXACT TEMPLATE WITH EXCELLENT STRUCTURE):
{subject}

Request for approval regarding [specific details from summary]. [Brief context or objective sentence]."""

    # Add bullet points only if user requested them
    if need_bullets:
        prompt += """

• [Generate 3 unique, specific bullet points based on the subject and summary - each should be 1 sentence, highly relevant to the specific event/request, not generic]
• [Focus on key aspects like: event details, participants, objectives, requirements, or benefits specific to this request]
• [Make each bullet point distinct and valuable, directly related to the subject matter]"""

    prompt += f"""

CRITICAL REQUIREMENTS FOR EXCELLENT CONTENT:
1. First line: ONLY "{subject}" (no prefixes like "Subject:")
2. Empty line
3. Request paragraph: EXACTLY 2-3 SHORT sentences starting with "Request for approval regarding" (VERY CONCISE - maximum 3 sentences total)
4. Empty line"""
    
    if need_bullets:
        prompt += """
5. Exactly 3 bullet points with • symbol (1 sentence each, EXCELLENT and HIGHLY SPECIFIC to the subject and summary - NO generic content)
6. Each bullet point must be unique and directly related to the specific event/request described in the summary
7. Focus on concrete details like: specific participants, exact objectives, particular requirements, or specific benefits
8. Avoid generic phrases like "key requirements" or "important details" - be specific and meaningful"""
    else:
        prompt += """
5. NO bullet points - continue directly to conclusion"""
    
    prompt += f"""
9. Use ONLY user's summary content - NO generic text
10. VERY CONCISE for STRICT single page limit - KEEP STARTING PARAGRAPH TO MAXIMUM 3 SENTENCES TOTAL
11. No conclusion paragraph (will be added separately)
12. No markdown formatting
13. Make content EXCELLENT and specific to user's summary only
14. Generate EXCELLENT content based on user inputs
15. Make it professional and specific to the request
16. MAXIMUM CONCISENESS - starting paragraph MUST be maximum 3 sentences total - very brief
17. EXCELLENT STRUCTURE - well-organized but SHORT paragraphs and bullet points
18. PERFECT ALIGNMENT with user's subject and summary
19. For bullet points: Extract specific details from the summary like names, dates, locations, objectives, participants, or unique aspects of the request
20. Make bullet points actionable and informative, not just descriptive
21. KEEP SENTENCES SHORT AND DIRECT - avoid lengthy explanations

Generate ULTRA-CONCISE, EXCELLENT, specific content based ONLY on the user's summary. No generic content. Single page limit is MANDATORY. Make it EXCELLENT.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an EXCELLENT professional NFA writer who creates very concise, single-page documents with PERFECT structure. Always generate EXCELLENT, specific content based ONLY on user inputs with VERY SHORT starting paragraphs (maximum 3 sentences total). For bullet points, extract specific details from the summary like names, dates, locations, objectives, participants, or unique aspects of the request. Make bullet points actionable and informative, not generic. Create well-structured content with very brief starting paragraphs, bullet points, and conclusions that perfectly align with the user's subject and summary. MAXIMUM conciseness required - KEEP STARTING PARAGRAPH TO MAXIMUM 3 SENTENCES TOTAL."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=150,  # Increased slightly for better bullet point generation
            temperature=0.1  # Lower for consistency and conciseness
        )
        
        ai_content = response.choices[0].message.content.strip()
        
        # Ensure proper formatting
        lines = ai_content.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                formatted_lines.append(line)
        
        # Join with proper spacing
        formatted_content = '\n\n'.join(formatted_lines)
        
        # Add the conclusion based on NFA type
        if nfa_type == "advance":
            conclusion = "The above proposal is submitted for approval, and the advance amount may kindly be released to the organizing committee to conduct the event smoothly."
        else:  # default = reimbursement
            conclusion = "The above proposal is submitted for approval, and the amount may kindly be reimbursed to the organizing committee after the event upon submission of the online report, receipts, and GST bills."
        
        # Add the conclusion at the end
        final_content = f"{formatted_content}\n\n{conclusion}"
        
        print(f"✅ AI generated content following strict template: {final_content[:200]}...", file=sys.stderr)
        
        return final_content
        
    except Exception as e:
        print(f"⚠️ AI Error: {e}", file=sys.stderr)
        # Create ultra-concise fallback content following strict template
        if need_bullets:
            fallback_content = f"""{subject}

Request for approval regarding {summary}. This proposal requires administrative approval for successful execution.

• Key requirements must be met for approval
• Important details will be outlined  
• Financial details provided in table"""
        else:
            fallback_content = f"""{subject}

Request for approval regarding {summary}. This proposal requires administrative approval for successful execution."""
        
        # Add conclusion based on NFA type
        if nfa_type == "advance":
            conclusion = "The above proposal is submitted for approval, and the advance amount may kindly be released to the organizing committee to conduct the event smoothly."
        else:  # default = reimbursement
            conclusion = "The above proposal is submitted for approval, and the amount may kindly be reimbursed to the organizing committee after the event upon submission of the online report, receipts, and GST bills."
        
        final_fallback = f"{fallback_content}\n\n{conclusion}"
        return final_fallback

# ==========================
# Document Creation Functions
# ==========================
def clean_text_content(text):
    """Clean text content to prevent DOCX corruption"""
    try:
        if not text:
            return ""
        
        # Remove null characters and other problematic characters that cause XML corruption
        cleaned = text.replace('\x00', '').replace('\r', '').replace('\x01', '').replace('\x02', '')
        
        # Normalize line endings
        cleaned = cleaned.replace('\n', '\n').replace('\r\n', '\n')
        
        # Remove excessive whitespace
        cleaned = ' '.join(cleaned.split())
        
        return cleaned
        
    except Exception as e:
        print(f"Error cleaning text content: {e}", file=sys.stderr)
        return text

def create_proper_nfa_document(subject_line, body_text, closing_line, table_data, nfa_type):
    """Create a properly structured NFA document that matches the reference image exactly"""
    try:
        print("📝 Creating properly structured NFA document...", file=sys.stderr)
        
        # Create new document
        doc = Document()
        
        # Set page margins - optimized for single page
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Add header image if it exists - CRITICAL for proper document structure
        if header_image_path and os.path.exists(header_image_path):
            try:
                page_width = Inches(8.5) - Inches(0.75) - Inches(0.75)
                doc.add_picture(header_image_path, width=page_width)
                print(f"✅ Header image added: {header_image_path}", file=sys.stderr)
                
                # Add spacing after header
                doc.add_paragraph()
                doc.add_paragraph()
            except Exception as e:
                print(f"⚠️ Could not add header image: {e}", file=sys.stderr)
        else:
            print("⚠️ Header image not found, creating document without header", file=sys.stderr)
        
        # Add date - right aligned (as shown in the image)
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Date: {datetime.now().strftime('%d-%m-%Y')}")
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        date_run.font.size = Pt(11)
        date_run.font.name = 'Arial'
        
        # Add empty line
        doc.add_paragraph()
        
        # Add title - centered (EXACTLY as in reference)
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("Note For Approval (NFA)")
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run.bold = True
        title_run.font.size = Pt(12)
        title_run.font.name = 'Arial'
        
        # Add empty line
        doc.add_paragraph()
        
        # Add subject - justified (EXACTLY as in reference)
        subject_para = doc.add_paragraph()
        subject_label = subject_para.add_run("Subject: ")
        subject_label.bold = True
        subject_label.font.size = Pt(11)
        subject_label.font.name = 'Arial'
        
        subject_text = subject_para.add_run(clean_text_content(subject_line))
        subject_text.font.size = Pt(11)
        subject_text.font.name = 'Arial'
        
        subject_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        # Add empty line
        doc.add_paragraph()
        
        # Add body content - process each section (EXACTLY as in reference)
        if body_text:
            sections = [section.strip() for section in body_text.split('\n\n') if section.strip()]
            
            for section in sections:
                if '•' in section:
                    # This is bullet points section
                    lines = [line.strip() for line in section.split('\n') if line.strip()]
                    for line in lines:
                        if line.startswith('•'):
                            bullet_text = line[1:].strip()
                            if bullet_text:
                                bullet_para = doc.add_paragraph()
                                bullet_run = bullet_para.add_run(f"• {clean_text_content(bullet_text)}")
                                bullet_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                bullet_run.font.size = Pt(11)
                                bullet_run.font.name = 'Arial'
                else:
                    # Regular paragraph
                    if section:
                        para = doc.add_paragraph()
                        run = para.add_run(clean_text_content(section))
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        run.font.size = Pt(11)
                        run.font.name = 'Arial'
        
        # Add empty line
        doc.add_paragraph()
        
        # Add table if provided (EXACTLY as in reference with borders)
        if table_data and len(table_data) > 0:
            add_proper_table_to_document(doc, table_data)
            
            # Add empty line after table
            doc.add_paragraph()
            
            # Add conclusion after table - justified (EXACTLY as in reference)
            conclusion_para = doc.add_paragraph()
            conclusion_run = conclusion_para.add_run(clean_text_content(closing_line))
            conclusion_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            conclusion_run.font.size = Pt(11)
            conclusion_run.font.name = 'Arial'
        else:
            # No table - add conclusion directly after body content
            conclusion_para = doc.add_paragraph()
            conclusion_run = conclusion_para.add_run(clean_text_content(closing_line))
            conclusion_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            conclusion_run.font.size = Pt(11)
            conclusion_run.font.name = 'Arial'
        
        # Add empty line before signatures
        doc.add_paragraph()
        
        # Add signature layout (EXACTLY as in reference - 2x2 grid)
        add_proper_signature_layout(doc)
        
        print("✅ Properly structured NFA document created", file=sys.stderr)
        return doc
        
    except Exception as e:
        print(f"❌ Error creating proper NFA document: {e}", file=sys.stderr)
        raise e

def add_proper_table_to_document(doc, table_data):
    """Add table data to document with proper formatting matching reference image"""
    try:
        print(f"add_proper_table_to_document called with: {table_data}", file=sys.stderr)
        if not table_data or len(table_data) == 0:
            print("No table data to add", file=sys.stderr)
            return
        
        print(f"Adding table with {len(table_data)} rows", file=sys.stderr)
        
        # Create table
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 0)
        
        # Apply styling with borders (EXACTLY as in reference)
        table.style = 'Table Grid'
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        print(f"Created table with {len(table.rows)} rows and {len(table.columns)} columns", file=sys.stderr)
        
        # Add data to table
        for row_idx, row_data in enumerate(table_data):
            print(f"Adding row {row_idx}: {row_data}", file=sys.stderr)
            if row_idx < len(table.rows):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(table.rows[row_idx].cells):
                        # Clean and validate cell data
                        clean_data = str(cell_data).strip() if cell_data is not None else ""
                        clean_data = clean_data.replace('\x00', '').replace('\r', '').replace('\n', ' ')
                        
                        # Get the cell and its paragraph
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = clean_data
                        
                        # Format cell content (EXACTLY as in reference)
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            
                            # Format runs in the paragraph
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                                run.font.name = 'Arial'
                                if row_idx == 0:  # Header row - bold
                                    run.bold = True
                                    run.font.size = Pt(11)
                                elif row_idx == len(table_data) - 1:  # Total row - bold
                                    run.bold = True
                                    run.font.size = Pt(11)
        
        print(f"✅ Proper table added to document with {len(table_data)} rows", file=sys.stderr)
        
    except Exception as e:
        print(f"Error adding proper table to document: {e}", file=sys.stderr)

def add_table_to_document(doc, table_data):
    """Legacy function - redirects to proper table function"""
    add_proper_table_to_document(doc, table_data)

def add_proper_signature_layout(doc):
    """Add signature layout to document with proper 2x2 grid matching reference image exactly"""
    try:
        print("Starting to add proper signature layout", file=sys.stderr)
        
        # Add minimal space before signatures
        doc.add_paragraph()
        
        # Create first signature table (top row) with 3 columns for spacing
        table = doc.add_table(rows=4, cols=3)
        
        # Set column widths: left signature, spacer, right signature
        try:
            table.columns[0].width = Inches(2.2)  # Left signature
            table.columns[1].width = Inches(1.6)   # Spacer column
            table.columns[2].width = Inches(2.2)  # Right signature
        except Exception as e:
            print(f"Warning: Could not set column widths: {e}", file=sys.stderr)
        
        # Top row - signature lines
        top_left_cell = table.cell(0, 0)
        top_right_cell = table.cell(0, 2)
        
        top_left_cell.text = "_________________"
        top_right_cell.text = "_________________"
        
        # Second row - names
        name_left_cell = table.cell(1, 0)
        name_right_cell = table.cell(1, 2)
        
        name_left_cell.text = "Dr Phani Kumar Pullela"
        name_right_cell.text = "Mr Chandrasekhar KN"
        
        # Third row - designations
        desig_left_cell = table.cell(2, 0)
        desig_right_cell = table.cell(2, 2)
        
        desig_left_cell.text = "Dean, Student Affairs"
        desig_right_cell.text = "Head Finance"
        
        
        # Apply formatting to all cells (EXACTLY as in reference)
        for cell in [top_left_cell, top_right_cell, name_left_cell, name_right_cell, 
                     desig_left_cell, desig_right_cell]:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
        
        print("First signature table created successfully", file=sys.stderr)
        
        # Add proper vertical spacing between signature blocks
        doc.add_paragraph()
        
        # Second signature table (bottom row) with 3 columns for spacing
        table2 = doc.add_table(rows=4, cols=3)
        
        # Set column widths
        table2.columns[0].width = Inches(2.2)
        table2.columns[1].width = Inches(1.6)
        table2.columns[2].width = Inches(2.2)
        
        # Top row - signature lines
        top_left_cell2 = table2.cell(0, 0)
        top_right_cell2 = table2.cell(0, 2)
        
        top_left_cell2.text = "_________________"
        top_right_cell2.text = "_________________"
        
        # Second row - names
        name_left_cell2 = table2.cell(1, 0)
        name_right_cell2 = table2.cell(1, 2)
        
        name_left_cell2.text = "Dr Sahana D Gowda"
        name_right_cell2.text = "Prof (Dr) Dwarika Prasad Uniyal"
        
        # Third row - designations
        desig_left_cell2 = table2.cell(2, 0)
        desig_right_cell2 = table2.cell(2, 2)
        
        desig_left_cell2.text = "Registrar - RV University"
        desig_right_cell2.text = "Vice Chancellor (i/c)"
        
        # Apply formatting to all cells (EXACTLY as in reference)
        for cell in [top_left_cell2, top_right_cell2, name_left_cell2, name_right_cell2, 
                     desig_left_cell2, desig_right_cell2]:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
        
        print("✅ Both proper signature tables created successfully", file=sys.stderr)
        
    except Exception as e:
        print(f"❌ Error adding proper signature layout: {e}", file=sys.stderr)
        raise e

def add_signature_layout(doc):
    """Legacy function - redirects to proper signature function"""
    add_proper_signature_layout(doc)

def parse_table_data(table_data_json):
    """Parse table data from JSON string"""
    try:
        import json
        print(f"Raw table data JSON: {table_data_json}", file=sys.stderr)
        table_data = json.loads(table_data_json)
        print(f"Parsed table data: {table_data}", file=sys.stderr)
        if isinstance(table_data, list) and len(table_data) > 0:
            print(f"Table data is valid list with {len(table_data)} rows", file=sys.stderr)
            return table_data
        print("Table data is empty or not a list", file=sys.stderr)
        return []
    except Exception as e:
        print(f"Error parsing table data: {e}", file=sys.stderr)
        return []

def extract_document_text(doc):
    """Extract text content from document for preview"""
    try:
        text_content = []
        
        # Add header info
        text_content.append("RV UNIVERSITY")
        text_content.append("Go, change the world")
        text_content.append("an initiative of RV EDUCATIONAL INSTITUTIONS")
        text_content.append("")
        text_content.append("RV Vidyaniketan, 8th Mile, Mysuru Road, Bengaluru, 560059, India")
        text_content.append("Ph: +91 80 68199900 | www.rvu.edu.in")
        text_content.append("")
        text_content.append("Date: " + datetime.now().strftime("%d-%m-%Y"))
        text_content.append("")
        text_content.append("Note For Approval (NFA)")
        text_content.append("")
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract table content
        for table in doc.tables:
            text_content.append("")  # Add spacing before table
            text_content.append("Financial/Resource Implications Table:")
            text_content.append("")  # Add spacing before table content
            
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    # Format table rows with proper spacing and borders
                    if row_idx == 0:  # Header row
                        header_line = " | ".join(f" {text:^12} " for text in row_text)
                        text_content.append(header_line)
                        text_content.append("+" + "+".join("-" * 14 for _ in row_text) + "+")
                    else:
                        data_line = " | ".join(f" {text:^12} " for text in row_text)
                        text_content.append(data_line)
            
            text_content.append("")  # Add spacing after table
        
        return "\n".join(text_content)
        
    except Exception as e:
        print(f"Error extracting document text: {e}", file=sys.stderr)
        return "Error extracting document content"

# ==========================
# AI Edit Function
# ==========================
def process_ai_edit(original_text, edit_prompt):
    """Process AI edit request on existing NFA content"""
    if not client:
        return f"{original_text}\n\n[AI Edit Applied: {edit_prompt}]"
    
    try:
        # Create AI prompt for editing
        prompt = f"""
You are an expert NFA (Note For Approval) editor. You have been given an existing NFA document and a modification request.

EXISTING NFA CONTENT:
{original_text}

MODIFICATION REQUEST:
{edit_prompt}

CRITICAL REQUIREMENTS:
1. Apply ONLY the requested modification - do not change anything else
2. Maintain the EXACT original structure and format
3. Keep the same subject line, conclusion, and signature sections unchanged
4. Preserve all justified alignment and STRICT single page format
5. Ensure the document remains professional and formal
6. Do not regenerate the entire document - only apply the specific requested changes
7. Keep content ULTRA-CONCISE to maintain single page limit
8. Return the complete modified NFA document with minimal changes
9. Preserve the strict template structure:
   - Subject line (no "Subject:" prefix)
   - Request paragraph starting with "Request for approval regarding"
   - Bullet points (if present) with • symbol
   - Conclusion paragraph
   - Signature block
10. SINGLE PAGE LIMIT IS MANDATORY - every word must count
11. EXCELLENT STRUCTURE - maintain well-organized paragraphs and bullet points
12. PERFECT ALIGNMENT - ensure content perfectly aligns with user's subject and summary

IMPORTANT: Only modify what the user specifically requested. Keep everything else exactly the same. Maintain ULTRA-CONCISE format with EXCELLENCE.

Return the complete modified NFA document with the requested changes applied.
"""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an EXCELLENT professional NFA editor who applies specific modifications to existing documents while maintaining their PERFECT structure and format. Always keep changes minimal and preserve single page format. Create EXCELLENT, well-structured content that perfectly aligns with the user's subject and summary. ULTRA-CONCISE editing with EXCELLENCE required."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=200,  # Reduced for single page edits
            temperature=0.1  # Lower for consistency and conciseness
        )
        
        edited_content = response.choices[0].message.content.strip()
        
        print(f"✅ AI edit completed: {edit_prompt}", file=sys.stderr)
        print(f"✅ Edited content length: {len(edited_content)} characters", file=sys.stderr)
        
        return edited_content
        
    except Exception as e:
        print(f"⚠️ AI Edit Error: {e}", file=sys.stderr)
        # Return original text with edit note
        return f"{original_text}\n\n[AI Edit Applied: {edit_prompt}]"

# ==========================
# Generate DOCX from Edited Text
# ==========================
def generate_docx_from_text(edited_text, subject, summary, nfa_type, table_data=None):
    """Generate DOCX document from edited text content with original formatting"""
    try:
        print(f"generate_docx_from_text called with subject: {subject}", file=sys.stderr)
        
        # Create output directory if it doesn't exist - FIXED to match server static serving
        output_dir = os.path.join(os.path.dirname(__file__), "..", "generated_letters", "nfa")
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"edited_nfa_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)
        
        print(f"Creating document: {filepath}", file=sys.stderr)
        
        # Parse the edited text to extract components
        sections = [section.strip() for section in edited_text.split('\n\n') if section.strip()]
        
        # Extract subject (first section)
        subject_line = sections[0] if sections else subject
        if subject_line.lower().startswith("subject:"):
            subject_line = subject_line.split(":", 1)[1].strip()
        
        # Extract conclusion (last section that contains conclusion keywords)
        closing_line = "The above proposal is submitted for approval."
        for section in reversed(sections):
            if any(keyword in section.lower() for keyword in ["proposal is submitted", "request your approval", "kindly be released", "kindly be reimbursed"]):
                closing_line = section
                break
        
        # Extract body content (everything between subject and conclusion)
        body_sections = []
        for i, section in enumerate(sections):
            if i == 0:  # Skip subject
                continue
            if section == closing_line:  # Skip conclusion
                continue
            body_sections.append(section)
        
        body_text = "\n\n".join(body_sections).strip()
        
        # If body is empty or malformed, create fallback
        if not body_text or len(body_sections) < 1:
            print("⚠️ Body content is malformed, creating fallback", file=sys.stderr)
            body_text = f"Request for approval regarding {summary}. This proposal requires administrative approval."
        
        # Create properly structured document
        doc = create_proper_nfa_document(subject_line, body_text, closing_line, table_data, nfa_type)
        
        # Save document
        doc.save(filepath)
        
        print(f"✅ DOCX generated successfully: {filepath}", file=sys.stderr)
        print(f"📁 Full file path: {filepath}", file=sys.stderr)
        print(f"📄 File exists: {os.path.exists(filepath)}", file=sys.stderr)
        print(f"📊 File size: {os.path.getsize(filepath) if os.path.exists(filepath) else 'N/A'} bytes", file=sys.stderr)
        
        return f"/generated_letters/nfa/{filename}", filename
        
    except Exception as e:
        print(f"❌ Error generating DOCX: {e}", file=sys.stderr)
        raise e

# ==========================
# Main Function
# ==========================
def main():
    # Check if this is download mode
    if len(sys.argv) > 1 and sys.argv[1] == "--download-mode":
        if len(sys.argv) < 6:
            print(json.dumps({
                "success": False,
                "error": "Download mode requires: --download-mode, editedText, subject, summary, nfaType"
            }))
            sys.exit(1)
        
        edited_text = sys.argv[2]
        subject = sys.argv[3]
        summary = sys.argv[4]
        nfa_type = sys.argv[5]
        table_data_json = sys.argv[6] if len(sys.argv) > 6 else "[]"
        
        # Parse table data
        table_data = parse_table_data(table_data_json)
        
        # Generate DOCX from edited text
        try:
            file_path, file_name = generate_docx_from_text(edited_text, subject, summary, nfa_type, table_data)
            print(json.dumps({
                "success": True,
                "filePath": file_path,
                "fileName": file_name,
                "message": "Edited NFA document generated successfully"
            }))
        except Exception as e:
            print(json.dumps({
                "success": False,
                "error": f"Failed to generate DOCX: {str(e)}"
            }))
        return
    
    # Check if this is edit mode
    if len(sys.argv) > 1 and sys.argv[1] == "--edit-mode":
        if len(sys.argv) < 4:
            print(json.dumps({
                "success": False,
                "error": "Edit mode requires: --edit-mode, text, prompt"
            }))
            sys.exit(1)
        
        original_text = sys.argv[2]
        edit_prompt = sys.argv[3]
        
        # Process AI edit
        edited_text = process_ai_edit(original_text, edit_prompt)
        
        print(json.dumps({
            "success": True,
            "editedText": edited_text,
            "message": "NFA text edited successfully"
        }))
        return
    
    # Arguments from Node.js for normal generation
    if len(sys.argv) < 4:
        print("Usage: generate_nfa_automation.py <subject> <summary> <nfa_type> [bullets] [table_data]", file=sys.stderr)
        sys.exit(1)

    subject = sys.argv[1]
    summary = sys.argv[2]
    nfa_type = sys.argv[3].lower()
    need_bullets = sys.argv[4].lower() in ("yes", "y", "true", "1") if len(sys.argv) > 4 else False
    table_data_json = sys.argv[5] if len(sys.argv) > 5 else "[]"

    print(f"Inputs -> Subject: {subject}, Type: {nfa_type}, Bullets: {need_bullets}", file=sys.stderr)

    # Parse table data
    table_data = parse_table_data(table_data_json)
    print(f"Table data parsed: {len(table_data)} rows", file=sys.stderr)

    # Generate NFA with structured format
    nfa_text = generate_ai_nfa_from_summary(subject, summary, nfa_type, need_bullets=need_bullets, facts_only=False)

    # Parse AI output for structured format
    print(f"Raw AI output: {nfa_text[:300]}...", file=sys.stderr)
    
    # Split by double newlines to get sections
    sections = [section.strip() for section in nfa_text.split('\n\n') if section.strip()]
    
    print(f"Parsed sections: {len(sections)}", file=sys.stderr)
    for i, section in enumerate(sections):
        print(f"Section {i}: {section[:100]}...", file=sys.stderr)
    
    # Extract subject (first section)
    subject_line = sections[0] if sections else subject
    if subject_line.lower().startswith("subject:"):
        subject_line = subject_line.split(":", 1)[1].strip()
    
    # Extract conclusion (last section that contains conclusion keywords)
    closing_line = "The above proposal is submitted for approval."
    for section in reversed(sections):
        if any(keyword in section.lower() for keyword in ["proposal is submitted", "request your approval", "kindly be released", "kindly be reimbursed"]):
            closing_line = section
            break
    
    # Extract body content (everything between subject and conclusion)
    body_sections = []
    for i, section in enumerate(sections):
        if i == 0:  # Skip subject
            continue
        if section == closing_line:  # Skip conclusion
            continue
        body_sections.append(section)
    
    body_text = "\n\n".join(body_sections).strip()
    
    # If body is empty or malformed, create fallback
    if not body_text or len(body_sections) < 1:
        print("⚠️ Body content is malformed, creating fallback", file=sys.stderr)
        if need_bullets:
            body_text = f"Request for approval regarding {summary}. This proposal requires administrative approval.\n\n• Key requirements must be met for approval\n• Important details will be outlined\n• Financial details provided in table"
        else:
            body_text = f"Request for approval regarding {summary}. This proposal requires administrative approval."
    
    print(f"✅ Structured content parsed - Subject: {subject_line}", file=sys.stderr)
    print(f"✅ Body sections: {len(body_sections)}", file=sys.stderr)
    print(f"✅ Closing line: {closing_line}", file=sys.stderr)
    print(f"✅ Body text length: {len(body_text)} characters", file=sys.stderr)

    # Filename
    sanitized_subject = re.sub(r'[\\/:*?"<>|]', '_', subject_line.replace(' ', '_')[:60])
    filename = os.path.join(output_directory, f"NFA_{nfa_type}_{sanitized_subject}.docx")

    print(f"Creating document: {filename}", file=sys.stderr)

    # Create properly structured document using the new function
    try:
        print("📝 Creating properly structured NFA document...", file=sys.stderr)
        doc = create_proper_nfa_document(subject_line, body_text, closing_line, table_data, nfa_type)
        print("✅ Properly structured document created successfully", file=sys.stderr)
    except Exception as e:
        print(f"❌ Error creating properly structured document: {e}", file=sys.stderr)
        raise

    # Save the properly structured document
    try:
        # Save document with error handling
        doc.save(filename)
        print(f"Document saved successfully: {filename}", file=sys.stderr)
        
        # Verify file was created and is readable
        if os.path.exists(filename):
            file_size = os.path.getsize(filename)
            print(f"File created successfully, size: {file_size} bytes", file=sys.stderr)
            if file_size == 0:
                print("Warning: File is empty!", file=sys.stderr)
        else:
            print("Error: File was not created!", file=sys.stderr)
            raise Exception("Document file was not created")
        
        # Extract text content for preview
        nfa_text_content = extract_document_text(doc)
        print(f"NFA text content extracted: {len(nfa_text_content)} characters", file=sys.stderr)
        
    except Exception as e:
        print(f"Error saving document: {e}", file=sys.stderr)
        print(f"Document validation failed for file: {filename}", file=sys.stderr)
        raise

    # Print relative path and text content for Node.js
    relative_path = os.path.relpath(filename, backend_dir)
    
    # Output structured JSON result
    result = {
        "success": True,
        "file_path": relative_path,
        "nfa_text": nfa_text_content,
        "file_name": os.path.basename(filename)
    }
    print(json.dumps(result))

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        error_result = {
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }
        print(json.dumps(error_result))
        print(f"Script failed: {e}", file=sys.stderr)
        sys.exit(1)
