import sys
import os
import re
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv

# Check Python version
print(f"Python version: {sys.version}", file=sys.stderr)
print(f"Python executable: {sys.executable}", file=sys.stderr)

# Load env vars and initialize OpenAI client
try:
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("Warning: OPENAI_API_KEY not found in environment variables", file=sys.stderr)
    else:
        print("OPENAI_API_KEY loaded successfully", file=sys.stderr)
    
    # Initialize OpenAI client with new API
    client = OpenAI(api_key=api_key)
    print("OpenAI client initialized successfully", file=sys.stderr)
except Exception as e:
    print(f"Error loading environment variables: {e}", file=sys.stderr)
    client = None

# Paths (from uploads folder)
script_dir = os.path.dirname(os.path.abspath(__file__))
backend_dir = os.path.dirname(script_dir)
uploads_dir = os.path.join(backend_dir, "uploads")

header_image_path = os.path.join(uploads_dir, "header.png")

# Output directory - organize into ms_reco folder
env_output_dir = os.getenv("OUTPUT_DIR")
if env_output_dir:
    # If OUTPUT_DIR is set in .env, resolve it relative to backend directory
    if env_output_dir.startswith("./backend/"):
        # Remove the "./backend/" prefix since we're already in the backend directory
        env_output_dir = env_output_dir.replace("./backend/", "./")
    
    # Resolve the path relative to backend directory
    base_output_directory = os.path.abspath(os.path.join(backend_dir, env_output_dir.lstrip("./")))
else:
    # Default fallback
    base_output_directory = os.path.join(uploads_dir, "generated_letters")

# Create ms_reco subfolder
output_directory = os.path.join(base_output_directory, "ms_reco")

print(f"Final output directory: {output_directory}", file=sys.stderr)
print(f"Environment OUTPUT_DIR: {env_output_dir}", file=sys.stderr)
print(f"Backend directory: {backend_dir}", file=sys.stderr)
os.makedirs(output_directory, exist_ok=True)

def generate_letter(name, title1, para1):
    return f"""
I am writing to wholeheartedly recommend {name} for admission to your esteemed Master's program. Having had the privilege of working closely with {name} during their academic journey at RV University, I can confidently attest to their exceptional academic abilities, strong work ethic, and potential for success in graduate studies.

Their significant contributions to "{title1}" have been commendable. {name} demonstrated remarkable dedication, analytical thinking, and problem-solving skills throughout the project. Their ability to work independently while also contributing effectively to team efforts has been impressive.

{para1}

In summary, {name} is a student of exceptional talent and dedication who would be a valuable addition to your Master's program. I have no doubt that they will excel in their graduate studies and make significant contributions to their field.

Sincerely,
Dr. Phani Kumar Pullela
Professor & Associate Dean
RV University
"""

def generate_ai_paragraph(name, title, summary):
    if not client:
        return f"{name} has demonstrated exceptional skills and dedication in their project '{title}'. {summary}"
    
    prompt = f"""
Write a concise and professional paragraph about {name}'s project titled '{title}'.
The project summary is: {summary}.
Highlight strengths, contributions, and skills that would be relevant for graduate school admission.
Focus on academic excellence, research potential, and leadership qualities.
Keep it to 3-4 sentences maximum.
"""
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a professional academic letter generator specializing in graduate school recommendations."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=100,
            temperature=0.7
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"Warning: AI generation failed, using fallback text: {e}", file=sys.stderr)
        return f"{name} has demonstrated exceptional skills and dedication in their project '{title}'. {summary}"

def main():
    # Read arguments from Node.js (form inputs)
    name = sys.argv[1]
    title1 = sys.argv[2]
    summary1 = sys.argv[3]

    # Debug: Print paths and working directory to stderr (not captured by Node.js)
    print(f"Current working directory: {os.getcwd()}", file=sys.stderr)
    print(f"Script directory: {script_dir}", file=sys.stderr)
    print(f"Backend directory: {backend_dir}", file=sys.stderr)
    print(f"Uploads directory: {uploads_dir}", file=sys.stderr)
    print(f"Header image path: {header_image_path}", file=sys.stderr)
    print(f"Output directory: {output_directory}", file=sys.stderr)
    
    # Check if header image exists
    if os.path.exists(header_image_path):
        print(f"OK: header.png found at: {header_image_path}", file=sys.stderr)
    else:
        print(f"ERROR: header.png NOT found at: {header_image_path}", file=sys.stderr)

    # Check OpenAI client status
    if client:
        print("OpenAI client ready for AI generation", file=sys.stderr)
    else:
        print("Warning: OpenAI client not available, will use fallback text", file=sys.stderr)

    ai_paragraph = generate_ai_paragraph(name, title1, summary1)
    letter = generate_letter(name, title1, ai_paragraph)

    # Save as docx
    sanitized_name = re.sub(r'[\\/:*?"<>|]', '_', name.replace(' ', '_'))
    filename = os.path.join(output_directory, f"MS_Recommendation_Letter_{sanitized_name}.docx")
    
    print(f"Creating document: {filename}", file=sys.stderr)

    try:
        doc = Document()
        
        # Add header image if it exists
        if os.path.exists(header_image_path):
            doc.add_picture(header_image_path, width=Inches(6.0))
            print("Header image added successfully", file=sys.stderr)
        else:
            print("Warning: Header image not found, skipping image addition", file=sys.stderr)
            
    except Exception as e:
        print(f"Error adding header image: {e}", file=sys.stderr)
        print(f"   Header image path: {header_image_path}", file=sys.stderr)
        # Continue without image if it fails
        doc = Document()
    
    # Add content
    p = doc.add_paragraph("Letter of Recommendation")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph("""
Dr. Phani Kumar Pullela
Professor & Associate Dean
RV University
Email: phanikumarp@rvu.edu.in
Phone: +91-9902005868
    """)
    
    subject_paragraph = doc.add_paragraph()
    subject_paragraph.add_run("Subject:").bold = True
    subject_paragraph.add_run(f" Recommendation Letter for {name}")
    
    t = doc.add_paragraph("To Whomsoever It May Concern")
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    p1 = doc.add_paragraph(letter)
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Set font and size for main content
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    # Make title bold
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    
    try:
        doc.save(filename)
        print(f"Document saved successfully: {filename}", file=sys.stderr)
    except Exception as e:
        print(f"Error saving document: {e}", file=sys.stderr)
        print(f"   Output directory: {output_directory}", file=sys.stderr)
        print(f"   Filename: {filename}", file=sys.stderr)
        raise

    # Print only the relative path for Node.js (not the full absolute path)
    relative_path = os.path.relpath(filename, backend_dir)
    print(relative_path)  # This will be: uploads/generated_letters/ms_reco/MS_Recommendation_Letter_*.docx

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"Script failed with error: {e}", file=sys.stderr)
        sys.exit(1)
