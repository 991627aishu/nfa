import sys
import os
import re
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv

# Check Python version
print(f"Python version: {sys.version}", file=sys.stderr)
print(f"Python executable: {sys.executable}", file=sys.stderr)

# Load env vars and initialize OpenAI client
try:
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("Warning: OPENAI_API_KEY not found in environment variables", file=sys.stderr)
    else:
        print("OPENAI_API_KEY loaded successfully", file=sys.stderr)
    
    # Initialize OpenAI client with new API
    client = OpenAI(api_key=api_key)
    print("OpenAI client initialized successfully", file=sys.stderr)
except Exception as e:
    print(f"Error loading environment variables: {e}", file=sys.stderr)
    client = None

# Paths (from uploads folder)
script_dir = os.path.dirname(os.path.abspath(__file__))
backend_dir = os.path.dirname(script_dir)
uploads_dir = os.path.join(backend_dir, "uploads")

header_image_path = os.path.join(uploads_dir, "header.png")

# Output directory - organize into job_reco folder
env_output_dir = os.getenv("OUTPUT_DIR")
if env_output_dir:
    # If OUTPUT_DIR is set in .env, resolve it relative to backend directory
    if env_output_dir.startswith("./backend/"):
        # Remove the "./backend/" prefix since we're already in the backend directory
        env_output_dir = env_output_dir.replace("./backend/", "./")
    
    # Resolve the path relative to backend directory
    base_output_directory = os.path.abspath(os.path.join(backend_dir, env_output_dir.lstrip("./")))
else:
    # Default fallback
    base_output_directory = os.path.join(uploads_dir, "generated_letters")

# Create job_reco subfolder
output_directory = os.path.join(base_output_directory, "job_reco")

print(f"Final output directory: {output_directory}", file=sys.stderr)
print(f"Environment OUTPUT_DIR: {env_output_dir}", file=sys.stderr)
print(f"Backend directory: {backend_dir}", file=sys.stderr)
os.makedirs(output_directory, exist_ok=True)

def generate_job_letter(name, title1, para1):
    return f"""
It is my privilege to recommend {name}, who has consistently demonstrated dedication and expertise in their work. I have had the pleasure of mentoring {name}, and their significant contributions to "{title1}" have been truly commendable.

{para1}

In addition to their technical expertise, {name} has shown exceptional professionalism, teamwork, and the ability to manage complex tasks with precision. Their commitment to excellence makes them a standout individual who will undoubtedly excel in any professional endeavor.

I am confident that {name} will bring the same level of passion, innovation, and excellence to their future roles. Should you require additional information, please do not hesitate to contact me at +91-9902005868 or via email at phanikumarp@rvu.edu.in.

Sincerely,
"""

def generate_ai_paragraph(name, title, summary):
    if not client:
        return f"{name} has demonstrated exceptional skills and dedication in their project '{title}'. {summary}"
    
    prompt = f"""
Write a very concise and professional paragraph (maximum 3-4 sentences) about {name}'s project titled '{title}'.
The project summary is: {summary}.
Focus on highlighting their key strengths and contributions. Keep it brief and impactful for a single-page recommendation letter.
"""
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a professional letter generator specializing in job recommendations. Keep responses very concise for single-page letters."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=100,  # Reduced for brevity
            temperature=0.7
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"Warning: AI generation failed, using fallback text: {e}", file=sys.stderr)
        return f"{name} has demonstrated exceptional skills and dedication in their project '{title}'. {summary}"

def main():
    # Read arguments from Node.js (form inputs)
    name = sys.argv[1]
    title1 = sys.argv[2]
    summary1 = sys.argv[3]

    # Debug: Print paths and working directory to stderr (not captured by Node.js)
    print(f"Current working directory: {os.getcwd()}", file=sys.stderr)
    print(f"Script directory: {script_dir}", file=sys.stderr)
    print(f"Backend directory: {backend_dir}", file=sys.stderr)
    print(f"Uploads directory: {uploads_dir}", file=sys.stderr)
    print(f"Header image path: {header_image_path}", file=sys.stderr)
    print(f"Output directory: {output_directory}", file=sys.stderr)
    
    # Check if header image exists
    if os.path.exists(header_image_path):
        print(f"OK: header.png found at: {header_image_path}", file=sys.stderr)
    else:
        print(f"ERROR: header.png NOT found at: {header_image_path}", file=sys.stderr)

    # Check OpenAI client status
    if client:
        print("OpenAI client ready for AI generation", file=sys.stderr)
    else:
        print("Warning: OpenAI client not available, will use fallback text", file=sys.stderr)

    ai_paragraph = generate_ai_paragraph(name, title1, summary1)
    letter = generate_job_letter(name, title1, ai_paragraph)

    # Save as docx
    sanitized_name = re.sub(r'[\\/:*?"<>|]', '_', name.replace(' ', '_'))
    filename = os.path.join(output_directory, f"Job_Recommendation_Letter_{sanitized_name}.docx")
    
    print(f"Creating document: {filename}", file=sys.stderr)

    try:
        doc = Document()
        
        # Set margins for single page optimization
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # Add header image if it exists
        if os.path.exists(header_image_path):
            # Calculate full page width (page width - margins)
            page_width = Inches(8.5) - Inches(0.7) - Inches(0.7)  # Full page minus margins
            doc.add_picture(header_image_path, width=page_width)
            print("Header image added successfully", file=sys.stderr)
        else:
            print("Header image not found, skipping...", file=sys.stderr)
    except Exception as e:
        print(f"Error adding header image: {e}", file=sys.stderr)
        print(f"   Header image path: {header_image_path}", file=sys.stderr)
        print("Continuing without header image...", file=sys.stderr)
    
    # Add date
    from datetime import datetime
    date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    p = doc.add_paragraph("Letter of Recommendation")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph("""
Dr. Phani Kumar Pullela
Professor & Associate Dean
RV University
Email: phanikumarp@rvu.edu.in
Phone: +91-9902005868
    """)
    
    subject_paragraph = doc.add_paragraph()
    subject_paragraph.add_run("Subject:").bold = True
    subject_paragraph.add_run(f" Job Recommendation Letter for {name}")
    
    t = doc.add_paragraph("To Whomsoever It May Concern")
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    p1 = doc.add_paragraph(letter)
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Add signature section (text only, no image)
    p2 = doc.add_paragraph("""Dr. Phani Kumar Pullela
Professor & Associate Dean
RV University""")
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Set font styling for single page optimization
    p1.style.font.name = 'Times New Roman'
    p1.style.font.size = Pt(11)  # Reduced from 12 to 11
    p2.style.font.name = 'Times New Roman'
    p2.style.font.size = Pt(11)  # Reduced from 12 to 11
    
    # Set paragraph spacing for compact layout
    p1.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.space_after = Pt(6)
    
    try:
        doc.save(filename)
        print(f"Document saved successfully: {filename}", file=sys.stderr)
    except Exception as e:
        print(f"Error saving document: {e}", file=sys.stderr)
        print(f"   Output directory: {output_directory}", file=sys.stderr)
        print(f"   Filename: {filename}", file=sys.stderr)
        raise

    # Print only the relative path for Node.js (not the full absolute path)
    relative_path = os.path.relpath(filename, backend_dir)
    print(relative_path)  # This will be: uploads/generated_letters/Job_Recommendation_Letter_*.docx

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"Script failed with error: {e}", file=sys.stderr)
        sys.exit(1)
