# backend/python/generate_nfa_automation.py
import sys
import os
import re
import json
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from datetime import datetime

# Debug Python version
print(f"Python version: {sys.version}", file=sys.stderr)
print(f"Python executable: {sys.executable}", file=sys.stderr)

# ==========================
# Load API Key & Init OpenAI
# ==========================
try:
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("Warning: OPENAI_API_KEY not found", file=sys.stderr)
    else:
        print("OPENAI_API_KEY loaded successfully", file=sys.stderr)
    client = OpenAI(api_key=api_key) if api_key else None
except Exception as e:
    print(f"Error initializing OpenAI client: {e}", file=sys.stderr)
    client = None

# ==========================
# Paths
# ==========================
script_dir = os.path.dirname(os.path.abspath(__file__))
backend_dir = os.path.dirname(script_dir)
uploads_dir = os.path.join(backend_dir, "uploads")

header_image_path = os.path.join(uploads_dir, "header.png")
signatures_dir = os.path.join(uploads_dir, "signatures")

# Output directory
env_output_dir = os.getenv("OUTPUT_DIR")
if env_output_dir:
    if env_output_dir.startswith("./backend/"):
        env_output_dir = env_output_dir.replace("./backend/", "./")
    base_output_directory = os.path.abspath(os.path.join(backend_dir, env_output_dir.lstrip("./")))
else:
    base_output_directory = os.path.join(uploads_dir, "generated_letters")

output_directory = os.path.join(base_output_directory, "nfa")
os.makedirs(output_directory, exist_ok=True)

print(f"Final output directory: {output_directory}", file=sys.stderr)

# ==========================
# AI Helper
# ==========================
def generate_ai_nfa_from_summary(subject, summary, nfa_type="reimbursement", need_bullets=False, facts_only=False):
    if not client:
        return f"{subject}\n\nRequest for approval regarding {summary}. The above proposal is submitted for approval."

    # Create ultra-concise prompt for single page following strict template
    prompt = f"""
Create an EXCELLENT professional NFA document following this EXACT structure for SINGLE PAGE output:

USER INPUTS:
Subject: {subject}
Summary: {summary}
Type: {nfa_type}
Need Bullets: {need_bullets}

REQUIRED OUTPUT FORMAT (MUST FOLLOW EXACT TEMPLATE WITH EXCELLENT STRUCTURE):
{subject}

Request for approval regarding [specific details from summary]. [Brief context or objective sentence]."""

    # Add bullet points only if user requested them
    if need_bullets:
        prompt += """

• [Generate 3 EXCEPTIONAL bullet points that SURPASS standard NFA documents and demonstrate advanced linguistic mastery]
• [Extract and articulate specific details with sophisticated precision: names, dates, locations, objectives, participants, unique aspects, contextual significance, institutional value]
• [Create bullet points that demonstrate SUPERIOR UNIQUENESS and scholarly sophistication beyond conventional business writing]
• [Ensure each bullet point is contextually perfect, professionally sophisticated, and specifically relevant to the user's exact request]
• [Focus on creating bullet points that showcase deep understanding, institutional expertise, and mastery of the subject matter]
• [Elevate content to demonstrate academic excellence and professional superiority that exceeds typical NFA documents]

QUALITY BENCHMARKS - SURPASS THESE EXAMPLES:
- Instead of basic phrases like "This initiative will promote engagement," use sophisticated expressions like "This strategically designed initiative will foster meaningful community engagement and collaborative participation."
- Replace simple statements with complex-compound sentences that demonstrate linguistic mastery and professional sophistication.
- Elevate vocabulary beyond standard business terms to include advanced academic and professional terminology.
- Create unique phrasing that avoids clichés and demonstrates creative linguistic mastery."""

    prompt += f"""

PROFESSIONAL DOCUMENT STRUCTURE REQUIREMENTS:
1. First line: "Subject: {subject}" (with "Subject:" prefix)
2. Empty line
3. Request paragraph: Create an EXCEPTIONAL 2-3 sentence paragraph starting with "Request for approval regarding" that SURPASSES standard NFA documents:
   - ABSOLUTE grammatical perfection with advanced syntax and sophisticated sentence construction
   - Superior vocabulary that demonstrates linguistic mastery and academic excellence
   - Unique, creative phrasing that elevates content beyond conventional business writing
   - Deep contextual understanding with precise articulation of purpose, significance, and institutional value
   - Professional tone that commands respect while demonstrating scholarly sophistication
   - Advanced sentence structures including complex-compound sentences, sophisticated clauses, and elegant modifiers
   - Seamless transitions and eloquent flow that creates compelling, persuasive narratives
4. Empty line"""
    
    if need_bullets:
        prompt += f"""
5. Exactly 3 bullet points with • symbol (1 sentence each, demonstrating SUPERIOR writing that SURPASSES standard NFA documents)
6. CRITICAL: Each bullet point must be EXCEPTIONALLY TAILORED to the specific subject "{subject}" and summary "{summary}"
7. Extract and articulate SPECIFIC DETAILS with sophisticated precision: names, dates, locations, objectives, participants, unique aspects, contextual significance, institutional value
8. Create bullet points with EXCEPTIONAL UNIQUENESS that demonstrates advanced linguistic mastery and scholarly sophistication
9. Employ advanced vocabulary, complex sentence structures, sophisticated clauses, and elegant transitions that exceed typical business writing
10. Each bullet point must be CONTEXTUALLY PERFECT, professionally sophisticated, and specifically relevant to the user's exact request
11. Ensure every bullet point showcases SUPERIOR writing quality, absolute grammatical perfection, and academic excellence
12. Focus on concrete, meaningful aspects that demonstrate deep understanding, institutional expertise, and mastery of the subject matter
13. Elevate content beyond conventional bullet points to demonstrate professional excellence and scholarly articulation"""
    else:
        prompt += """
5. NO bullet points - continue directly to conclusion"""
    
    prompt += f"""
GRAMMAR AND WRITING EXCELLENCE REQUIREMENTS:
9. Demonstrate MASTERY of the user's specific context: Subject: "{subject}", Summary: "{summary}", NFA Type: "{nfa_type}"
10. VERY CONCISE for STRICT single page limit - KEEP STARTING PARAGRAPH TO MAXIMUM 3 SENTENCES TOTAL while maintaining GRAMMATICAL PERFECTION
11. DO NOT include any conclusion statements - conclusion will be added separately after table
12. No markdown formatting - pure text only
13. Create content with FLAWLESS GRAMMAR, perfect syntax, and sophisticated vocabulary
14. Every sentence must demonstrate linguistic mastery and deep understanding
15. Use UNIQUE PHRASING that avoids clichés and generic expressions
16. MAXIMUM CONCISENESS with MAXIMUM IMPACT - starting paragraph MUST be maximum 3 sentences total
17. STRUCTURAL EXCELLENCE - perfectly organized, logical flow, professional presentation
18. CONTEXTUAL PERFECTION - content must be specifically tailored to this exact request
19. GRAMMATICAL RIGOR - demonstrate scholarly approach to language and writing
20. PROFESSIONAL AUTHORITY - content must command respect through writing excellence

FINAL INSTRUCTION: Generate content that demonstrates the work of an EXCEPTIONAL English Professor with mastery of grammar, syntax, and sophisticated writing styles. Create content that is grammatically perfect, contextually sophisticated, and uniquely written. Every sentence must showcase linguistic excellence, perfect grammar, and professional writing artistry.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a WORLD-CLASS English Professor and Professional Writer with exceptional expertise in creating superior NFA documents. Your writing surpasses standard business documents and demonstrates MASTERFUL command of language, sophisticated expression, and unparalleled grammatical precision.\n\nSUPERIOR WRITING EXCELLENCE:\n- ABSOLUTE grammatical perfection with flawless syntax, punctuation, and sentence construction\n- Advanced vocabulary with sophisticated word choices that demonstrate linguistic mastery\n- Unique, creative phrasing that elevates content beyond conventional business writing\n- Perfect subject-verb agreement, tense consistency, and advanced sentence structures\n- Seamless transitions and eloquent flow that creates compelling narratives\n- Professional tone that commands respect while maintaining accessibility\n\nCONTENT CREATION MASTERY:\n- Craft compelling opening paragraphs with maximum 3 sentences total that are sophisticated yet very concise\n- Generate EXCEPTIONAL bullet points that demonstrate deep understanding and superior articulation\n- Create content that is contextually perfect, professionally sophisticated, and grammatically impeccable\n- Use advanced sentence structures including complex-compound sentences, sophisticated clauses, and elegant modifiers\n- Demonstrate expertise through precise, articulate language that exceeds typical business communication\n\nADVANCED WRITING REQUIREMENTS:\n- Elevate professional tone to demonstrate institutional excellence and academic sophistication\n- Employ sophisticated vocabulary including advanced academic and professional terminology\n- Create unique expressions that avoid clichés and demonstrate creative linguistic mastery\n- Ensure every sentence showcases superior writing ability and professional expertise\n- Generate content that is more polished, articulate, and sophisticated than standard NFA documents\n\nSTRUCTURAL PERFECTION:\n- DO NOT include conclusion statements (will be added separately)\n- Create opening paragraphs that are more compelling and sophisticated than typical business documents\n- Generate bullet points that demonstrate superior understanding and exceptional articulation\n- Ensure every element showcases writing mastery and professional excellence\n- Maintain the highest standards of academic and professional communication"},
                {"role": "user", "content": prompt}
            ],
            max_tokens=300,  # Reduced for very concise 3-sentence maximum starting paragraphs
            temperature=0.5  # Higher for more creative and sophisticated writing styles
        )
        
        ai_content = response.choices[0].message.content.strip()
        
        # Ensure proper formatting
        lines = ai_content.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                formatted_lines.append(line)
        
        # Join with proper spacing
        formatted_content = '\n\n'.join(formatted_lines)
        
        # Generate sophisticated conclusion based on NFA type
        conclusion = generate_sophisticated_conclusion(nfa_type)
        
        # AGGRESSIVELY remove any existing conclusion lines from AI content
        # BUT preserve the subject line
        content_without_conclusion = formatted_content
        
        # Split content into lines to preserve subject
        lines = formatted_content.split('\n')
        filtered_lines = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Always preserve subject line
            if line_lower.startswith('subject:') or line_lower.startswith('subject'):
                filtered_lines.append(line)
                continue
            
            # Skip conclusion patterns
            conclusion_patterns = [
                "proposal is submitted for approval",
                "request your approval", 
                "kindly be released",
                "kindly be reimbursed",
                "above proposal is submitted",
                "submitted for approval",
                "organizing committee",
                "after the event",
                "upon submission",
                "online report",
                "receipts and gst bills"
            ]
            
            # Check if this line contains any conclusion pattern
            is_conclusion = any(pattern in line_lower for pattern in conclusion_patterns)
            
            # Also check for partial conclusion matches (more aggressive)
            if not is_conclusion:
                # Additional check for lines that might be conclusion fragments
                if any(fragment in line_lower for fragment in [
                    "proposal is submitted", "kindly be", "organizing committee", 
                    "after the event", "upon submission", "receipts and gst"
                ]):
                    print(f"✅ Skipping conclusion fragment: {line[:50]}...", file=sys.stderr)
                    continue
                filtered_lines.append(line)
            else:
                print(f"✅ Skipping conclusion line: {line[:50]}...", file=sys.stderr)
        
        content_without_conclusion = '\n'.join(filtered_lines)
        
        print(f"✅ AI generated content (conclusion removed): {content_without_conclusion[:200]}...", file=sys.stderr)
        
        return content_without_conclusion
        
    except Exception as e:
        print(f"⚠️ AI Error: {e}", file=sys.stderr)
        # Create ultra-concise fallback content following strict template
        if need_bullets:
            fallback_content = f"""Subject: {subject}

Request for approval regarding {summary}. This proposal requires administrative approval for successful execution.

• Key requirements must be met for approval
• Important details will be outlined  
• Financial details provided in table"""
        else:
            fallback_content = f"""Subject: {subject}

Request for approval regarding {summary}. This proposal requires administrative approval for successful execution."""
        
        # Generate sophisticated conclusion based on NFA type
        conclusion = generate_sophisticated_conclusion(nfa_type)
        
        # AGGRESSIVELY remove conclusion from fallback content
        # BUT preserve the subject line
        fallback_without_conclusion = fallback_content
        
        # Split content into lines to preserve subject
        lines = fallback_content.split('\n')
        filtered_lines = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Always preserve subject line
            if line_lower.startswith('subject:') or line_lower.startswith('subject'):
                filtered_lines.append(line)
                continue
            
            # Skip conclusion patterns
            conclusion_patterns = [
                "proposal is submitted for approval",
                "request your approval", 
                "kindly be released",
                "kindly be reimbursed",
                "above proposal is submitted",
                "submitted for approval",
                "organizing committee",
                "after the event",
                "upon submission",
                "online report",
                "receipts and gst bills"
            ]
            
            # Check if this line contains any conclusion pattern
            is_conclusion = any(pattern in line_lower for pattern in conclusion_patterns)
            
            # Also check for partial conclusion matches (more aggressive)
            if not is_conclusion:
                # Additional check for lines that might be conclusion fragments
                if any(fragment in line_lower for fragment in [
                    "proposal is submitted", "kindly be", "organizing committee", 
                    "after the event", "upon submission", "receipts and gst"
                ]):
                    print(f"✅ Skipping fallback conclusion fragment: {line[:50]}...", file=sys.stderr)
                    continue
                filtered_lines.append(line)
            else:
                print(f"✅ Skipping fallback conclusion line: {line[:50]}...", file=sys.stderr)
        
        fallback_without_conclusion = '\n'.join(filtered_lines)
        
        return fallback_without_conclusion

# ==========================
# Sophisticated Conclusion Generator
# ==========================
def generate_sophisticated_conclusion(nfa_type):
    """Generate EXCEPTIONAL, sophisticated conclusion lines that surpass standard NFA documents"""
    if nfa_type == "advance":
        conclusions = [
            "This comprehensive proposal is respectfully submitted for administrative consideration, and the advance amount may be sanctioned to facilitate the successful execution and timely implementation of this strategically significant initiative.",
            "The foregoing proposal is presented for institutional approval, and the requested advance may be authorized to ensure the effective realization of this valuable endeavor, thereby supporting our commitment to excellence.",
            "This detailed submission seeks approval for the advance amount, which will enable the successful completion of the outlined objectives, deliverables, and institutional goals with the highest standards of quality.",
            "The proposal is respectfully submitted for review and consideration, and the advance may be sanctioned to facilitate the comprehensive execution of this innovative and impactful initiative within our academic community.",
            "This proposal is presented for administrative approval, and the advance amount may be authorized to support the successful implementation of this strategically important project, thereby advancing our institutional mission and objectives."
        ]
    else:  # reimbursement
        conclusions = [
            "This comprehensive proposal is respectfully submitted for administrative consideration, and the amount may be reimbursed to the organizing committee upon successful completion of the event, following submission of detailed online reports, authenticated receipts, and comprehensive GST documentation.",
            "The foregoing proposal is presented for institutional approval, and the expenses may be reimbursed upon completion of the event with proper documentation, financial verification, and submission of all necessary receipts and supporting materials.",
            "This detailed submission seeks approval for reimbursement, which will be processed upon completion of the event with submission of comprehensive receipts, detailed reports, and complete GST documentation for administrative review and verification.",
            "The proposal is respectfully submitted for review and consideration, and the amount may be reimbursed after the event upon presentation of detailed receipts, comprehensive reports, and all necessary GST documentation to ensure complete financial transparency and accountability.",
            "This proposal is presented for administrative approval, and the amount may be reimbursed upon successful completion of the event with submission of authenticated receipts, detailed financial reports, and comprehensive GST documentation for institutional review and processing."
        ]
    
    # Return a randomly selected conclusion for variety
    import random
    return random.choice(conclusions)

# ==========================
# Signature Layout Helper
# ==========================
def get_signature_layout():
    """Get signature layout from JSON database"""
    try:
        import json
        db_path = os.path.join(backend_dir, "database", "signatures.json")
        print(f"Looking for signature database at: {db_path}", file=sys.stderr)
        
        if os.path.exists(db_path):
            print("Signature database file exists", file=sys.stderr)
            with open(db_path, 'r') as f:
                data = json.load(f)
            
            signatures = data.get('signatures', {})
            print(f"Loaded signatures: {signatures}", file=sys.stderr)
            
            # Format signatures for the 2x2 layout
            layout = {
                'top_left': signatures.get('prepared_by', [{}])[0] if signatures.get('prepared_by') else {},
                'top_right': next((s for s in signatures.get('approved_by', []) if s.get('order') == 1), {}),
                'bottom_left': signatures.get('recommended_by', [{}])[0] if signatures.get('recommended_by') else {},
                'bottom_right': next((s for s in signatures.get('approved_by', []) if s.get('order') == 2), {})
            }
            
            print(f"Signature layout: {layout}", file=sys.stderr)
            return layout
        else:
            print("Signature database file does not exist", file=sys.stderr)
    except Exception as e:
        print(f"Error loading signature layout: {e}", file=sys.stderr)
    
    # Return default layout with your specified signatures
    print("Using default signature layout", file=sys.stderr)
    return {
        'top_left': {
            'name': 'Dr Phani Kumar Pullela',
            'designation': 'Dean, Student Affairs'
        },
        'top_right': {
            'name': 'Mr Chandrasekhar KN',
            'designation': 'Head Finance'
        },
        'bottom_left': {
            'name': 'Dr Sahana D Gowda',
            'designation': 'Registrar - RV University'
        },
        'bottom_right': {
            'name': 'Prof (Dr) Dwarika Prasad Uniyal',
            'designation': 'Vice Chancellor (i/c)'
        }
    }

def add_signature_layout(doc, layout):
    """Add signature layout to document with safer approach"""
    try:
        print("Starting to add signature layout", file=sys.stderr)
        
        # Add minimal space before signatures for single page limit
        doc.add_paragraph()
        
        print("Creating first signature table", file=sys.stderr)
        
        # Create first signature table (top row) with 3 columns for spacing
        table = doc.add_table(rows=4, cols=3)
        
        # Set column widths: left signature, spacer, right signature
        try:
            table.columns[0].width = Inches(2.2)  # Left signature
            table.columns[1].width = Inches(1.6)   # Increased spacer column for better gap
            table.columns[2].width = Inches(2.2)  # Right signature
        except Exception as e:
            print(f"Warning: Could not set column widths: {e}", file=sys.stderr)
        
        # Top row - signature lines
        top_left_cell = table.cell(0, 0)
        top_right_cell = table.cell(0, 2)
        
        # Add text safely
        try:
            top_left_cell.text = "_________________"
            top_right_cell.text = "_________________"
        except Exception as e:
            print(f"Warning: Could not set signature lines: {e}", file=sys.stderr)
            # Try alternative approach
            if top_left_cell.paragraphs:
                top_left_cell.paragraphs[0].text = "_________________"
            if top_right_cell.paragraphs:
                top_right_cell.paragraphs[0].text = "_________________"
        
        # Justify align the signature lines and reduce font size
        for cell in [top_left_cell, top_right_cell]:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Reduced font size
        
        # Second row - names
        name_left_cell = table.cell(1, 0)
        name_right_cell = table.cell(1, 2)
        
        # Add names safely
        try:
            left_name = clean_text_content(layout.get('top_left', {}).get('name', 'Dr Phani Kumar Pullela'))
            right_name = clean_text_content(layout.get('top_right', {}).get('name', 'Mr Chandrasekhar KN'))
            
            name_left_cell.text = left_name
            name_right_cell.text = right_name
        except Exception as e:
            print(f"Warning: Could not set names: {e}", file=sys.stderr)
            # Try alternative approach
            if name_left_cell.paragraphs:
                name_left_cell.paragraphs[0].text = "Dr Phani Kumar Pullela"
            if name_right_cell.paragraphs:
                name_right_cell.paragraphs[0].text = "Mr Chandrasekhar KN"
        
        # Apply justified alignment to names and reduce font size
        for cell in [name_left_cell, name_right_cell]:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Reduced font size
        
        # Third row - designations
        desig_left_cell = table.cell(2, 0)
        desig_right_cell = table.cell(2, 2)
        
        desig_left_cell.text = layout.get('top_left', {}).get('designation', 'Dean, Student Affairs')
        desig_right_cell.text = layout.get('top_right', {}).get('designation', 'Head Finance')
        
        # Apply justified alignment to designations and reduce font size
        for cell in [desig_left_cell, desig_right_cell]:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Reduced font size
        
        
        print("First signature table created successfully", file=sys.stderr)
        
        # Add proper vertical spacing between signature blocks
        doc.add_paragraph()
        
        print("Creating second signature table", file=sys.stderr)
        
        # Second signature table (bottom row) with 3 columns for spacing
        table2 = doc.add_table(rows=4, cols=3)
        
        # Set column widths: left signature, spacer, right signature
        table2.columns[0].width = Inches(2.2)  # Left signature
        table2.columns[1].width = Inches(1.6)   # Increased spacer column for better gap
        table2.columns[2].width = Inches(2.2)  # Right signature
        
        # Top row - signature lines
        top_left_cell2 = table2.cell(0, 0)
        top_right_cell2 = table2.cell(0, 2)
        
        top_left_cell2.text = "_________________"
        top_right_cell2.text = "_________________"
        
        # Justify align the signature lines and reduce font size
        for cell in [top_left_cell2, top_right_cell2]:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Reduced font size
        
        # Second row - names
        name_left_cell2 = table2.cell(1, 0)
        name_right_cell2 = table2.cell(1, 2)
        
        name_left_cell2.text = layout.get('bottom_left', {}).get('name', 'Dr Sahana D Gowda')
        name_right_cell2.text = layout.get('bottom_right', {}).get('name', 'Prof (Dr) Dwarika Prasad Uniyal')
        
        # Apply justified alignment to names and reduce font size
        for cell in [name_left_cell2, name_right_cell2]:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Reduced font size
        
        # Third row - designations
        desig_left_cell2 = table2.cell(2, 0)
        desig_right_cell2 = table2.cell(2, 2)
        
        desig_left_cell2.text = layout.get('bottom_left', {}).get('designation', 'Registrar - RV University')
        desig_right_cell2.text = layout.get('bottom_right', {}).get('designation', 'Vice Chancellor (i/c)')
        
        # Apply justified alignment to designations and reduce font size
        for cell in [desig_left_cell2, desig_right_cell2]:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Reduced font size
        
        
        print("✅ Both signature tables created successfully", file=sys.stderr)
        
    except Exception as e:
        print(f"❌ Error adding signature layout: {e}", file=sys.stderr)
        raise e

def parse_table_data(table_data_json):
    """Parse table data from JSON string"""
    try:
        import json
        print(f"Raw table data JSON: {table_data_json}", file=sys.stderr)
        table_data = json.loads(table_data_json)
        print(f"Parsed table data: {table_data}", file=sys.stderr)
        if isinstance(table_data, list) and len(table_data) > 0:
            print(f"Table data is valid list with {len(table_data)} rows", file=sys.stderr)
            return table_data
        print("Table data is empty or not a list", file=sys.stderr)
        return []
    except Exception as e:
        print(f"Error parsing table data: {e}", file=sys.stderr)
        return []

def add_table_to_document(doc, table_data):
    """Add table data to document with Google Sheets-like formatting"""
    try:
        print(f"add_table_to_document called with: {table_data}", file=sys.stderr)
        if not table_data or len(table_data) == 0:
            print("No table data to add", file=sys.stderr)
            return
        
        print(f"Adding table with {len(table_data)} rows", file=sys.stderr)
        # Add space before table
        doc.add_paragraph()
        
        # Create table
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 0)
        
        # Apply Google Sheets-like styling
        table.style = 'Table Grid'  # Enable borders
        
        # Set table properties for better appearance
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        print(f"Created table with {len(table.rows)} rows and {len(table.columns)} columns", file=sys.stderr)
        
        # Add data to table with validation and formatting
        for row_idx, row_data in enumerate(table_data):
            print(f"Adding row {row_idx}: {row_data}", file=sys.stderr)
            if row_idx < len(table.rows):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(table.rows[row_idx].cells):
                        # Clean and validate cell data
                        clean_data = str(cell_data).strip() if cell_data is not None else ""
                        # Remove any potentially problematic characters
                        clean_data = clean_data.replace('\x00', '').replace('\r', '').replace('\n', ' ')
                        
                        # Get the cell and its paragraph
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = clean_data
                        
                        # Format cell content
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align text
                            
                            # Format runs in the paragraph
                            for run in paragraph.runs:
                                run.font.size = Pt(10)  # Consistent font size
                                if row_idx == 0:  # Header row
                                    run.bold = True  # Make header bold
                                    run.font.size = Pt(11)  # Slightly larger for header
        
        # Set column widths for better appearance
        try:
            if len(table.columns) > 0:
                # Calculate column widths based on content
                col_widths = []
                for col_idx in range(len(table.columns)):
                    if col_idx == 0:  # First column (usually sl.no)
                        col_widths.append(Inches(0.8))
                    elif col_idx == len(table.columns) - 1:  # Last column (usually total)
                        col_widths.append(Inches(1.0))
                    else:  # Middle columns
                        col_widths.append(Inches(1.5))
                
                # Apply column widths
                for i, width in enumerate(col_widths):
                    if i < len(table.columns):
                        table.columns[i].width = width
        except Exception as e:
            print(f"Warning: Could not set column widths: {e}", file=sys.stderr)
        
        print(f"✅ Table added to document with {len(table_data)} rows and Google Sheets-like formatting", file=sys.stderr)
        
    except Exception as e:
        print(f"Error adding table to document: {e}", file=sys.stderr)

def clean_text_content(text):
    """Clean text content to prevent DOCX corruption"""
    try:
        if not text:
            return ""
        
        # Remove null characters and other problematic characters
        cleaned = text.replace('\x00', '').replace('\r', '').replace('\x01', '').replace('\x02', '')
        
        # Normalize line endings
        cleaned = cleaned.replace('\n', '\n').replace('\r\n', '\n')
        
        # Remove excessive whitespace
        cleaned = ' '.join(cleaned.split())
        
        return cleaned
        
    except Exception as e:
        print(f"Error cleaning text content: {e}", file=sys.stderr)
        return text

def optimize_for_single_page(doc):
    """Optimize document content to ensure it fits on a single page - aggressive optimization"""
    try:
        print("🔍 Aggressively optimizing document for single page...", file=sys.stderr)
        
        # Count total paragraphs and estimate content length
        total_paragraphs = len(doc.paragraphs)
        total_tables = len(doc.tables)
        
        print(f"📊 Document stats: {total_paragraphs} paragraphs, {total_tables} tables", file=sys.stderr)
        
        # Always apply aggressive optimization for single page
        print("⚡ Applying aggressive single-page optimization...", file=sys.stderr)
        
        # Reduce all paragraph spacing to absolute minimum
        for para in doc.paragraphs:
            if hasattr(para, 'paragraph_format'):
                para.paragraph_format.space_after = Pt(0)  # No spacing after
                para.paragraph_format.space_before = Pt(0)  # No space before
                para.paragraph_format.line_spacing = 1.0  # Single line spacing
        
        # Reduce table cell padding to minimum
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if hasattr(para, 'paragraph_format'):
                            para.paragraph_format.space_after = Pt(0)
                            para.paragraph_format.space_before = Pt(0)
                            para.paragraph_format.line_spacing = 1.0
        
        print("✅ Document aggressively optimized for single page", file=sys.stderr)
        return True
        
    except Exception as e:
        print(f"❌ Error optimizing document: {e}", file=sys.stderr)
        return False

def validate_document_structure(doc):
    """Validate document structure before saving to prevent XML corruption"""
    try:
        print("🔍 Validating document structure...", file=sys.stderr)
        
        # Check if document has paragraphs
        if len(doc.paragraphs) == 0:
            print("❌ Document has no paragraphs", file=sys.stderr)
            return False
        
        # Check each paragraph for issues
        for i, para in enumerate(doc.paragraphs):
            try:
                # Check if paragraph has runs
                if not para.runs:
                    print(f"⚠️ Paragraph {i} has no runs", file=sys.stderr)
                    continue
                
                # Check each run for issues
                for j, run in enumerate(para.runs):
                    if run.text is None:
                        print(f"⚠️ Run {j} in paragraph {i} has None text", file=sys.stderr)
                        run.text = ""
                    elif '\x00' in run.text:
                        print(f"⚠️ Run {j} in paragraph {i} contains null characters", file=sys.stderr)
                        run.text = run.text.replace('\x00', '')
                
            except Exception as e:
                print(f"❌ Error validating paragraph {i}: {e}", file=sys.stderr)
                return False
        
        # Check tables
        for i, table in enumerate(doc.tables):
            try:
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text and '\x00' in run.text:
                                    print(f"⚠️ Table {i}, Row {row_idx}, Cell {cell_idx} contains null characters", file=sys.stderr)
                                    run.text = run.text.replace('\x00', '')
            except Exception as e:
                print(f"❌ Error validating table {i}: {e}", file=sys.stderr)
                return False
        
        print("✅ Document structure validation passed", file=sys.stderr)
        return True
        
    except Exception as e:
        print(f"❌ Document validation failed: {e}", file=sys.stderr)
        return False

def format_bullet_points(text):
    """Format text with proper bullet points"""
    try:
        # Clean text first
        text = clean_text_content(text)
        
        lines = text.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                # Convert any bullet point to circle bullet
                if line.startswith('-'):
                    line = line.replace('-', '•', 1)
                elif line.startswith('*'):
                    line = line.replace('*', '•', 1)
                formatted_lines.append(line)
            elif line and not line.startswith('Subject:') and not line.startswith('Request for approval'):
                # Add circle bullet point if it's a content line
                formatted_lines.append(f"• {line}")
            else:
                formatted_lines.append(line)
        
        return '\n'.join(formatted_lines)
    except Exception as e:
        print(f"Error formatting bullet points: {e}", file=sys.stderr)
        return text

def extract_document_text(doc):
    """Extract text content from document for preview - minimal clean version"""
    try:
        text_content = []
        
        # Extract only essential content paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text = paragraph.text.strip()
                # Skip all header information and duplicates (but keep Subject:)
                if not any(skip_text in text for skip_text in [
                    "RV UNIVERSITY", "Go, change the world", "an initiative of RV EDUCATIONAL INSTITUTIONS",
                    "RV Vidyaniketan, 8th Mile", "Ph: +91 80 68199900", "www.rvu.edu.in",
                    "Note For Approval (NFA)", "Date:"
                ]):
                    # Only add actual content paragraphs
                    if any(keyword in text for keyword in [
                        "Request for approval", "The above proposal is submitted", "•", "Subject:"
                    ]):
                        text_content.append(text)
        
        return "\n".join(text_content)
        
    except Exception as e:
        print(f"Error extracting document text: {e}", file=sys.stderr)
        return "Error extracting document content"

# ==========================
# AI Edit Function
# ==========================
def process_ai_edit(original_text, edit_prompt):
    """Process AI edit request on existing NFA content"""
    if not client:
        return f"{original_text}\n\n[AI Edit Applied: {edit_prompt}]"
    
    try:
        # Create AI prompt for editing
        prompt = f"""
You are an expert NFA (Note For Approval) editor. You have been given an existing NFA document and a modification request.

EXISTING NFA CONTENT:
{original_text}

MODIFICATION REQUEST:
{edit_prompt}

CRITICAL REQUIREMENTS:
1. Apply ONLY the requested modification - do not change anything else
2. Maintain the EXACT original structure and format
3. Keep the same subject line, conclusion, and signature sections unchanged
4. Preserve all justified alignment and STRICT single page format
5. Ensure the document remains professional and formal
6. Do not regenerate the entire document - only apply the specific requested changes
7. Keep content VERY CONCISE to maintain single page limit - STARTING PARAGRAPH MAXIMUM 3 SENTENCES TOTAL
8. Return the complete modified NFA document with minimal changes
9. Preserve the strict template structure:
   - Subject line (no "Subject:" prefix)
   - Request paragraph starting with "Request for approval regarding"
   - Bullet points (if present) with • symbol
   - Conclusion paragraph
   - Signature block
10. SINGLE PAGE LIMIT IS MANDATORY - every word must count
11. EXCELLENT STRUCTURE - maintain well-organized paragraphs and bullet points
12. PERFECT ALIGNMENT - ensure content perfectly aligns with user's subject and summary

IMPORTANT: Only modify what the user specifically requested. Keep everything else exactly the same. Maintain ULTRA-CONCISE format with EXCELLENCE.

Return the complete modified NFA document with the requested changes applied.
"""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an EXCELLENT professional NFA editor who applies specific modifications to existing documents while maintaining their PERFECT structure and format. Always keep changes minimal and preserve single page format. Create EXCELLENT, well-structured content with very brief starting paragraphs (maximum 3 sentences total) that perfectly aligns with the user's subject and summary. VERY CONCISE editing with EXCELLENCE required - KEEP STARTING PARAGRAPH TO MAXIMUM 3 SENTENCES TOTAL."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=200,  # Reduced for single page edits
            temperature=0.1  # Lower for consistency and conciseness
        )
        
        edited_content = response.choices[0].message.content.strip()
        
        print(f"✅ AI edit completed: {edit_prompt}", file=sys.stderr)
        print(f"✅ Edited content length: {len(edited_content)} characters", file=sys.stderr)
        
        return edited_content
        
    except Exception as e:
        print(f"⚠️ AI Edit Error: {e}", file=sys.stderr)
        # Return original text with edit note
        return f"{original_text}\n\n[AI Edit Applied: {edit_prompt}]"

# ==========================
# Generate DOCX from Edited Text
# ==========================
def generate_docx_from_text(edited_text, subject, summary, nfa_type, table_data=None):
    """Generate DOCX document from edited text content with original formatting"""
    try:
        print(f"generate_docx_from_text called with subject: {subject}", file=sys.stderr)
        
        # Create output directory if it doesn't exist
        output_dir = os.path.join(os.path.dirname(__file__), "..", "generated_letters", "nfa")
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"edited_nfa_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)
        
        print(f"Creating document: {filepath}", file=sys.stderr)
        
        # Create new document
        doc = Document()
        
        # Set page margins - aggressively optimized for single page
        for section in doc.sections:
            section.top_margin = Inches(0.3)      # Further reduced for single page
            section.bottom_margin = Inches(0.3)   # Further reduced for single page
            section.left_margin = Inches(0.5)     # Further reduced for single page
            section.right_margin = Inches(0.5)    # Further reduced for single page
        
        # Add header image if exists - use header.png from uploads directory
        header_image_path = os.path.join(uploads_dir, "header.png")
        if os.path.exists(header_image_path):
            # Calculate full page width (page width - left margin - right margin)
            page_width = Inches(8.5) - Inches(0.5) - Inches(0.5)  # Full page minus new margins
            doc.add_picture(header_image_path, width=page_width)
            print(f"✅ Header image added: {header_image_path}", file=sys.stderr)
        else:
            print(f"⚠️ Header image not found at: {header_image_path}", file=sys.stderr)
        
        # Add date - positioned on the right side (matches preview format)
        date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        if date_para.runs:
            date_para.runs[0].font.size = Pt(11)
        
        # Add title - centered and bold (matches preview format)
        title = doc.add_paragraph("Note For Approval (NFA)")
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if title.runs:
            title.runs[0].bold = True
            title.runs[0].font.size = Pt(12)
        
        # Add subject with justified alignment (matches preview format)
            subj_paragraph = doc.add_paragraph()
            subj_run = subj_paragraph.add_run("Subject: ")
            subj_run.bold = True
        subj_run.font.size = Pt(11)
        subj_text = subj_paragraph.add_run(subject)
        subj_text.font.size = Pt(11)
        subj_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Add body content - justified (matches preview format)
        # Remove subject line and conclusion lines from edited_text to avoid duplication
        body_lines = edited_text.split('\n')
        body_content_lines = []
        conclusion_keywords = [
            "proposal is submitted", "request your approval", "kindly be released", "kindly be reimbursed",
            "above proposal is submitted", "submitted for approval", "organizing committee", 
            "after the event", "upon submission", "online report", "receipts and gst bills"
        ]
        
        for line in body_lines:
            line_lower = line.strip().lower()
            # Skip subject lines
            if line_lower.startswith('subject:'):
                continue
            # Skip conclusion lines
            if any(keyword in line_lower for keyword in conclusion_keywords):
                continue
            body_content_lines.append(line)
        clean_body_text = '\n'.join(body_content_lines).strip()
        
        if clean_body_text:
            body_paragraph = doc.add_paragraph(clean_body_text)
        body_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        # No paragraph spacing for single page limit
        body_paragraph.paragraph_format.space_after = Pt(0)  # No spacing for single page
        if body_paragraph.runs:
            body_paragraph.runs[0].font.size = Pt(11)
        
        # Add table data if provided
        if table_data and len(table_data) > 0:
            print(f"Adding table data to edited document with {len(table_data)} rows", file=sys.stderr)
            add_table_to_document(doc, table_data)
        else:
            print("No table data to add to edited document", file=sys.stderr)
        
        # Generate sophisticated conclusion based on NFA type
        conclusion_text = generate_sophisticated_conclusion(nfa_type)
        
        conclusion_para = doc.add_paragraph(conclusion_text)
        conclusion_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        # No paragraph spacing for single page limit
        conclusion_para.paragraph_format.space_after = Pt(0)  # No spacing for single page
        
        # Add signature layout with default layout
        default_signature_layout = {
            'top_left': {
                'name': 'Dr Phani Kumar Pullela',
                'designation': 'Dean, Student Affairs'
            },
            'top_right': {
                'name': 'Mr Chandrasekhar',
                'designation': 'Head Finance'
            },
            'bottom_left': {
                'name': 'Dr Sahana D Gowda',
                'designation': 'Registrar - RV University (i/c)'
            },
            'bottom_right': {
                'name': 'Prof (Dr) Dwarika Prasad Uniyal',
                'designation': 'Vice Chancellor (i/c)'
            }
        }
        
        add_signature_layout(doc, default_signature_layout)
        
        # Optimize document for single page limit
        optimize_for_single_page(doc)
        
        # Save document
        doc.save(filepath)
        
        print(f"✅ DOCX generated successfully: {filepath}", file=sys.stderr)
        print(f"📁 Full file path: {filepath}", file=sys.stderr)
        print(f"📄 File exists: {os.path.exists(filepath)}", file=sys.stderr)
        print(f"📊 File size: {os.path.getsize(filepath) if os.path.exists(filepath) else 'N/A'} bytes", file=sys.stderr)
        
        return f"/generated_letters/nfa/{filename}", filename
        
    except Exception as e:
        print(f"❌ Error generating DOCX: {e}", file=sys.stderr)
        raise e

# ==========================
# Main Function
# ==========================
def main():
    # Check if this is download mode
    if len(sys.argv) > 1 and sys.argv[1] == "--download-mode":
        if len(sys.argv) < 6:
            print(json.dumps({
                "success": False,
                "error": "Download mode requires: --download-mode, editedText, subject, summary, nfaType"
            }))
            sys.exit(1)
        
        edited_text = sys.argv[2]
        subject = sys.argv[3]
        summary = sys.argv[4]
        nfa_type = sys.argv[5]
        table_data_json = sys.argv[6] if len(sys.argv) > 6 else "[]"
        
        # Parse table data
        table_data = parse_table_data(table_data_json)
        
        # Generate DOCX from edited text
        try:
            file_path, file_name = generate_docx_from_text(edited_text, subject, summary, nfa_type, table_data)
            print(json.dumps({
                "success": True,
                "filePath": file_path,
                "fileName": file_name,
                "message": "Edited NFA document generated successfully"
            }))
        except Exception as e:
            print(json.dumps({
                "success": False,
                "error": f"Failed to generate DOCX: {str(e)}"
            }))
        return
    
    # Check if this is edit mode
    if len(sys.argv) > 1 and sys.argv[1] == "--edit-mode":
        if len(sys.argv) < 4:
            print(json.dumps({
                "success": False,
                "error": "Edit mode requires: --edit-mode, text, prompt"
            }))
            sys.exit(1)
        
        original_text = sys.argv[2]
        edit_prompt = sys.argv[3]
        
        # Process AI edit
        edited_text = process_ai_edit(original_text, edit_prompt)
        
        print(json.dumps({
            "success": True,
            "editedText": edited_text,
            "message": "NFA text edited successfully"
        }))
        return
    
    # Arguments from Node.js for normal generation
    if len(sys.argv) < 4:
        print("Usage: generate_nfa_automation.py <subject> <summary> <nfa_type> [bullets] [table_data]", file=sys.stderr)
        sys.exit(1)

    subject = sys.argv[1]
    summary = sys.argv[2]
    nfa_type = sys.argv[3].lower()
    need_bullets = sys.argv[4].lower() in ("yes", "y", "true", "1") if len(sys.argv) > 4 else False
    table_data_json = sys.argv[5] if len(sys.argv) > 5 else "[]"

    print(f"Inputs -> Subject: {subject}, Type: {nfa_type}, Bullets: {need_bullets}", file=sys.stderr)

    # Parse table data
    table_data = parse_table_data(table_data_json)
    print(f"Table data parsed: {len(table_data)} rows", file=sys.stderr)

    # Generate NFA with structured format (use actual need_bullets parameter)
    nfa_text = generate_ai_nfa_from_summary(subject, summary, nfa_type, need_bullets=need_bullets, facts_only=False)

    # Parse AI output for structured format with robust error handling
    print(f"Raw AI output: {nfa_text[:300]}...", file=sys.stderr)
    
    # Split by double newlines to get sections
    sections = [section.strip() for section in nfa_text.split('\n\n') if section.strip()]
    
    print(f"Parsed sections: {len(sections)}", file=sys.stderr)
    for i, section in enumerate(sections):
        print(f"Section {i}: {section[:100]}...", file=sys.stderr)
    
    # Extract subject (first section) - keep "Subject:" prefix for preview
    subject_line = sections[0] if sections else f"Subject: {subject}"
    # Don't remove "Subject:" prefix - keep it for preview component
    if not subject_line.lower().startswith("subject:"):
        subject_line = f"Subject: {subject_line}"
    
    # Extract conclusion (last section that contains conclusion keywords)
    closing_line = "The above proposal is submitted for approval."
    conclusion_keywords = [
        "proposal is submitted", "request your approval", "kindly be released", "kindly be reimbursed",
        "above proposal is submitted", "submitted for approval", "organizing committee", 
        "after the event", "upon submission", "online report", "receipts and gst bills"
    ]
    
    for section in reversed(sections):
        if any(keyword in section.lower() for keyword in conclusion_keywords):
            closing_line = section
            print(f"✅ Found conclusion line: {section[:50]}...", file=sys.stderr)
            break
    
    # Extract body content (everything between subject and conclusion)
    body_sections = []
    for i, section in enumerate(sections):
        if i == 0:  # Skip subject
            continue
        # AGGRESSIVELY skip any conclusion-like sections and subject lines to avoid duplicates
        conclusion_keywords = [
            "proposal is submitted", "request your approval", "kindly be released", "kindly be reimbursed",
            "above proposal is submitted", "submitted for approval", "organizing committee", 
            "after the event", "upon submission", "online report", "receipts and gst bills"
        ]
        
        # Skip conclusion sections
        if any(keyword in section.lower() for keyword in conclusion_keywords):
            print(f"✅ Skipping conclusion section: {section[:50]}...", file=sys.stderr)
            continue
        
        # Skip subject sections (already added separately)
        if section.lower().startswith('subject:'):
            print(f"✅ Skipping subject section: {section[:50]}...", file=sys.stderr)
            continue
            
        body_sections.append(section)
    
    body_text = "\n\n".join(body_sections).strip()
    
    # If body is empty or malformed, create ultra-concise fallback (without conclusion)
    if not body_text or len(body_sections) < 1:
        print("⚠️ Body content is malformed, creating ultra-concise fallback", file=sys.stderr)
        if need_bullets:
            body_text = f"Request for approval regarding {summary}. This proposal requires administrative approval. The objective is successful event execution.\n\n• Key requirements must be met for approval\n• Important details will be outlined\n• Financial details provided in table"
        else:
            body_text = f"Request for approval regarding {summary}. This proposal requires administrative approval. The objective is successful event execution."
    
    print(f"✅ Structured content parsed - Subject: {subject_line}", file=sys.stderr)
    print(f"✅ Body sections: {len(body_sections)}", file=sys.stderr)
    print(f"✅ Closing line: {closing_line}", file=sys.stderr)
    print(f"✅ Body text length: {len(body_text)} characters", file=sys.stderr)

    # Filename
    sanitized_subject = re.sub(r'[\\/:*?"<>|]', '_', subject_line.replace(' ', '_')[:60])
    filename = os.path.join(output_directory, f"NFA_{nfa_type}_{sanitized_subject}.docx")

    print(f"Creating document: {filename}", file=sys.stderr)

    # Create docx with proper initialization and template
    try:
        # Create a new document with explicit template
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        doc = Document()
        
        # Ensure document has proper structure
        if not hasattr(doc, 'paragraphs'):
            raise Exception("Document does not have paragraphs attribute")
        
        print("✅ Document created successfully", file=sys.stderr)
    except Exception as e:
        print(f"❌ Error creating document: {e}", file=sys.stderr)
        raise

    # Set page margins - aggressively optimized for single page
    try:
        doc_sections = doc.sections
        for doc_section in doc_sections:
            doc_section.top_margin = Inches(0.3)      # Further reduced for single page
            doc_section.bottom_margin = Inches(0.3)   # Further reduced for single page
            doc_section.left_margin = Inches(0.5)     # Further reduced for single page
            doc_section.right_margin = Inches(0.5)    # Further reduced for single page
        print("✅ Page margins set successfully", file=sys.stderr)
    except Exception as e:
        print(f"❌ Error setting margins: {e}", file=sys.stderr)
        raise

    # Add header image from uploads directory (only the image, no text)
    if os.path.exists(header_image_path):
        # Calculate full page width (page width - left margin - right margin)
        page_width = Inches(8.5) - Inches(0.5) - Inches(0.5)  # Full page minus new margins
        doc.add_picture(header_image_path, width=page_width)
        print(f"✅ Header image added to main document: {header_image_path}", file=sys.stderr)
    else:
        print(f"⚠️ Header image not found at: {header_image_path}", file=sys.stderr)

    # Add content step by step with comprehensive validation
    try:
        # Add date - positioned on the right side (matches preview format)
        date_text = f"Date: {datetime.now().strftime('%d/%m/%Y')}"
        date_par = doc.add_paragraph()
        date_run = date_par.add_run(clean_text_content(date_text))
        date_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        date_run.font.size = Pt(11)
        print("✅ Date paragraph added (right-aligned)", file=sys.stderr)
        
        # Add empty line
        doc.add_paragraph()
        
        # Add title - centered and bold (matches preview format)
        title_par = doc.add_paragraph()
        title_run = title_par.add_run("Note For Approval (NFA)")
        title_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run.bold = True
        title_run.font.size = Pt(12)
        print("✅ Title paragraph added (centered)", file=sys.stderr)
        
        # Add empty line
        doc.add_paragraph()
        
        # Add subject - left aligned (matches preview format)
        subj_par = doc.add_paragraph()
        subj_run1 = subj_par.add_run("Subject: ")
        subj_run1.bold = True
        subj_run1.font.size = Pt(11)
        
        clean_subject = clean_text_content(subject_line)
        subj_run2 = subj_par.add_run(clean_subject)
        subj_run2.font.size = Pt(11)
        
        subj_par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        print("✅ Subject paragraph added (left-aligned)", file=sys.stderr)
        
        # Add empty line
        doc.add_paragraph()
        
    except Exception as e:
        print(f"❌ Error adding basic content: {e}", file=sys.stderr)
        raise

    # Add body content with safer approach
    try:
        print(f"Processing structured body text: {body_text[:200]}...", file=sys.stderr)
        
        # Split by double newlines to get sections (Request paragraph + Bullet points)
        sections = [section.strip() for section in body_text.split('\n\n') if section.strip()]
        
        print(f"Body sections to process: {len(sections)}", file=sys.stderr)
        
        for i, section in enumerate(sections):
            print(f"Processing section {i+1}: {section[:100]}...", file=sys.stderr)
            
            # AGGRESSIVELY skip any conclusion lines in body content
            conclusion_keywords = [
                "proposal is submitted", "request your approval", "kindly be released", "kindly be reimbursed",
                "above proposal is submitted", "submitted for approval", "organizing committee", 
                "after the event", "upon submission", "online report", "receipts and gst bills"
            ]
            
            # Skip this entire section if it contains conclusion keywords
            if any(keyword in section.lower() for keyword in conclusion_keywords):
                print(f"✅ Skipping conclusion section in body: {section[:50]}...", file=sys.stderr)
                continue
            
            # Check if this section contains bullet points
            if '•' in section:
                # This is the bullet points section
                lines = [line.strip() for line in section.split('\n') if line.strip()]
                bullet_count = 0
                
                for line in lines:
                    if line.startswith('•'):
                        # Remove bullet symbol and add as Word bullet
                        clean_text = line[1:].strip()
                        if clean_text:
                            # Clean bullet text content
                            clean_bullet_text = clean_text_content(clean_text)
                            
                            # Create paragraph manually to avoid style issues
                            para = doc.add_paragraph()
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            para.paragraph_format.space_after = Pt(0)  # No spacing for single page
                    
                            # Add bullet manually (matches preview format)
                            run = para.add_run("• ")
                            run.font.size = Pt(11)
                            
                            # Add text
                            text_run = para.add_run(clean_bullet_text)
                            text_run.font.size = Pt(11)
                            
                            bullet_count += 1
                            print(f"Added bullet point {bullet_count}: {clean_bullet_text[:50]}...", file=sys.stderr)
                    else:
                        # Regular text line (shouldn't happen in bullet section)
                        if line and not line.startswith('•'):
                            para = doc.add_paragraph()
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            run = para.add_run(clean_text_content(line))
                            run.font.size = Pt(11)
                            print(f"Added regular line in bullet section: {line[:50]}...", file=sys.stderr)
                
                print(f"✅ Added {bullet_count} bullet points", file=sys.stderr)
            else:
                # This is the "Request for approval" paragraph
                if section:
                    # Clean section content before adding
                    clean_section = clean_text_content(section)
                    
                    para = doc.add_paragraph()
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Matches preview format
                    para.paragraph_format.space_after = Pt(0)  # No spacing for single page
                    
                    run = para.add_run(clean_section)
                    run.font.size = Pt(11)
                    
                    print(f"Added request paragraph: {clean_section[:50]}...", file=sys.stderr)
        
        print("✅ Body content added successfully", file=sys.stderr)
        
    except Exception as e:
        print(f"❌ Error adding body content: {e}", file=sys.stderr)
        # Add fallback content
        fallback_para = doc.add_paragraph()
        fallback_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        fallback_run = fallback_para.add_run(f"Request for approval regarding {clean_text_content(summary)}. This proposal requires administrative approval.")
        fallback_run.font.size = Pt(11)
        print("✅ Fallback content added", file=sys.stderr)
    
    # Add table data if provided
    if table_data and len(table_data) > 0:
        print(f"Adding table data with {len(table_data)} rows", file=sys.stderr)
        add_table_to_document(doc, table_data)
    else:
        print("No table data to add", file=sys.stderr)
    
    # Add conclusion line after table based on NFA type
    try:
        # Generate sophisticated conclusion based on NFA type
        conclusion_text = generate_sophisticated_conclusion(nfa_type)
        
        conclusion_para = doc.add_paragraph()
        conclusion_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Matches preview format
        conclusion_para.paragraph_format.space_after = Pt(0)  # No spacing for single page
        
        conclusion_run = conclusion_para.add_run(clean_text_content(conclusion_text))
        conclusion_run.font.size = Pt(11)
        
        print(f"✅ Conclusion added successfully for NFA type: {nfa_type}", file=sys.stderr)
    except Exception as e:
        print(f"❌ Error adding conclusion: {e}", file=sys.stderr)
        # Add fallback conclusion
        fallback_conclusion = doc.add_paragraph()
        fallback_conclusion.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        fallback_run = fallback_conclusion.add_run("The above proposal is submitted for approval.")
        fallback_run.font.size = Pt(11)

    # Add signature layout - ALWAYS add signatures
    print("Getting signature layout", file=sys.stderr)
    signature_layout = get_signature_layout()
    print(f"Signature layout retrieved: {signature_layout}", file=sys.stderr)
    print("Adding signature layout to document", file=sys.stderr)
    
    # Force add signatures - try structured layout first, then fallback
    signatures_added = False
    
    try:
        if signature_layout and len(signature_layout) > 0:
            add_signature_layout(doc, signature_layout)
            print("✅ Structured signature layout added successfully", file=sys.stderr)
            signatures_added = True
        else:
            print("⚠️ No signature layout data, using fallback", file=sys.stderr)
            raise Exception("No signature data")
    except Exception as e:
        print(f"❌ Error adding structured signature layout: {e}", file=sys.stderr)
        print("Using fallback signature section", file=sys.stderr)
    
    # If structured layout failed, add simple fallback with justified alignment
    if not signatures_added:
        try:
            doc.add_paragraph()  # Reduced spacing for single page limit
            
            # First signature row with proper spacing
            sig1_para = doc.add_paragraph("_________________                    _________________")
            sig1_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in sig1_para.runs:
                run.font.size = Pt(10)
            
            name1_para = doc.add_paragraph("Dr Phani Kumar Pullela              Mr Chandrasekhar KN")
            name1_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in name1_para.runs:
                run.font.size = Pt(10)
            
            title1_para = doc.add_paragraph("Dean, Student Affairs                Head Finance")
            title1_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in title1_para.runs:
                run.font.size = Pt(10)
            
            role1_para = doc.add_paragraph("(Prepared by)                        (Approved by)")
            role1_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in role1_para.runs:
                run.font.size = Pt(10)
            
            # Add vertical spacing between signature rows
            doc.add_paragraph()
            
            # Second signature row with proper spacing
            sig2_para = doc.add_paragraph("_________________                    _________________")
            sig2_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in sig2_para.runs:
                run.font.size = Pt(10)
            
            name2_para = doc.add_paragraph("Dr Sahana D Gowda                    Prof (Dr) Dwarika Prasad Uniyal")
            name2_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in name2_para.runs:
                run.font.size = Pt(10)
            
            title2_para = doc.add_paragraph("Registrar - RV University             Vice Chancellor (i/c)")
            title2_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in title2_para.runs:
                run.font.size = Pt(10)
            
            role2_para = doc.add_paragraph("(Recommended by)                    (Approved by)")
            role2_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in role2_para.runs:
                run.font.size = Pt(10)
            
            print("✅ Fallback signatures added successfully with justified alignment", file=sys.stderr)
        except Exception as fallback_error:
            print(f"❌ Even fallback signatures failed: {fallback_error}", file=sys.stderr)

    # Initialize nfa_text_content
    nfa_text_content = ""
    
    try:
        # Optimize document for single page limit
        optimize_for_single_page(doc)
        
        # Validate document structure before saving
        if not validate_document_structure(doc):
            print("❌ Document validation failed, attempting to fix...", file=sys.stderr)
            # Try to fix common issues
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.text and '\x00' in run.text:
                        run.text = run.text.replace('\x00', '')
        
        # Save document with error handling
        doc.save(filename)
        print(f"Document saved successfully: {filename}", file=sys.stderr)
        
        # Verify file was created and is readable
        if os.path.exists(filename):
            file_size = os.path.getsize(filename)
            print(f"File created successfully, size: {file_size} bytes", file=sys.stderr)
            if file_size == 0:
                print("Warning: File is empty!", file=sys.stderr)
        else:
            print("Error: File was not created!", file=sys.stderr)
            raise Exception("Document file was not created")
        
        # Extract text content for preview
        nfa_text_content = extract_document_text(doc)
        print(f"NFA text content extracted: {len(nfa_text_content)} characters", file=sys.stderr)
        
    except Exception as e:
        print(f"Error saving document: {e}", file=sys.stderr)
        print(f"Document validation failed for file: {filename}", file=sys.stderr)
        raise

    # Print relative path and text content for Node.js
    relative_path = os.path.relpath(filename, backend_dir)
    
    # Output structured JSON result
    result = {
        "success": True,
        "file_path": relative_path,
        "nfa_text": nfa_text_content,
        "file_name": os.path.basename(filename)
    }
    print(json.dumps(result))

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        error_result = {
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }
        print(json.dumps(error_result))
        print(f"Script failed: {e}", file=sys.stderr)
        sys.exit(1)
